"""Coverage for Instrument's redesigned form reading (Task 12), its
per-form report artifact (Task 13), and the fixture sanity checks Task 11's
own verification calls for.

Follows the dependency-free convention already used by
test_schema_guardian.py / test_lexicon_librarian.py: plain
``def test_...()`` driving coroutines with ``asyncio.run(...)``, no live
LLM key, no Mongo, agents built directly with ``llm=None`` and a fake db.
"""
from __future__ import annotations

import asyncio
import json
from pathlib import Path
from shutil import which

import pytest
from pypdf import PdfReader

import phi_core.agents.specialists as specialists
from phi_core.agents.specialists import Instrument
from phi_core import paths as paths_mod

FIXTURES = Path(__file__).parent / "fixtures"
FLAT_PDF = FIXTURES / "tb_collection_form.pdf"
ACROFORM_PDF = FIXTURES / "tb_collection_form_acroform.pdf"
SCANNED_PDF = FIXTURES / "tb_collection_form_scanned.pdf"
GROUND_TRUTH = FIXTURES / "tb_collection_form_ground_truth.json"


# ---- shared fakes ----------------------------------------------------------


class FakeAgentLog:
    def __init__(self):
        self.inserted: list[dict] = []

    async def insert_one(self, doc, *_args, **_kwargs):
        self.inserted.append(doc)


class FakeDb:
    def __init__(self):
        self.agent_log = FakeAgentLog()


class RecordingInstrument(Instrument):
    """Never touches an LLM: records every call_json invocation (prompt +
    phase) and returns a caller-supplied canned reply."""

    def __init__(self, *args, reply=None, **kwargs):
        super().__init__(*args, **kwargs)
        self.call_json_calls: list[dict] = []
        self._reply = reply if reply is not None else {"fields": []}

    async def call_json(self, user_prompt, phase, default=None, **kwargs):
        self.call_json_calls.append({"user_prompt": user_prompt, "phase": phase})
        return self._reply


class NoLlmInstrument(Instrument):
    """Raises if Instrument ever reaches for an LLM -- proves a Tier-1
    AcroForm PDF costs zero LLM calls."""

    async def call_json(self, *args, **kwargs):
        raise AssertionError("Instrument must not call an LLM for a Tier-1 AcroForm PDF")


def _instrument(cls=RecordingInstrument, **kwargs) -> Instrument:
    return cls(session_id="s1", llm=None, db=FakeDb(), **kwargs)


def _form_file(file_id: str, path: Path, original_name: str | None = None) -> dict:
    return {
        "file_id": file_id,
        "stored_path": str(path),
        "original_name": original_name or path.name,
    }


def _ground_truth() -> dict:
    return json.loads(GROUND_TRUTH.read_text(encoding="utf-8"))


# ---- Task 11: fixture sanity / tier-routing smoke --------------------------


def test_fixtures_route_to_the_right_tier_by_get_fields():
    assert PdfReader(str(FLAT_PDF)).get_fields() is None
    assert PdfReader(str(SCANNED_PDF)).get_fields() is None
    acro_fields = PdfReader(str(ACROFORM_PDF)).get_fields()
    assert acro_fields and len(acro_fields) == len(_ground_truth()["fields"])


def test_scanned_fixture_digital_text_is_under_ocr_threshold():
    from phi_core.file_readers import OCR_TEXT_THRESHOLD

    reader = PdfReader(str(SCANNED_PDF))
    text = "\n\n".join((p.extract_text() or "") for p in reader.pages)
    assert len(text.strip()) < OCR_TEXT_THRESHOLD


# ---- Task 12: AcroForm Tier 1, zero LLM calls -------------------------------


def test_acroform_routes_to_tier1_with_zero_llm_calls():
    inst = _instrument(NoLlmInstrument)
    result = asyncio.run(inst.run(form_files=[_form_file("f1", ACROFORM_PDF)]))
    truth = _ground_truth()["fields"]
    assert len(result["fields"]) == len(truth)
    # Every AcroForm widget carries a real internal field name (unlike the
    # flat/scanned Tier-2 path, where an unannotated field's
    # collected_variable is null): the ground truth's annotated subset must
    # all be present, since those are the labels that were literally
    # printed in brackets next to the widget.
    variables = {f["collected_variable"] for f in result["fields"]}
    annotated_truth = {f["collected_variable"] for f in truth if f["collected_variable"]}
    assert annotated_truth <= variables
    assert inst._fields["f1"] == result["fields"]
    assert "phi_category" not in json.dumps(result)


# ---- Task 12: flat PDF Tier 2 through the shared reader ---------------------


def test_flat_form_routes_to_tier2_through_shared_read_pdf(monkeypatch):
    calls: list[Path] = []
    real_read_pdf = specialists.read_pdf

    def spy_read_pdf(path):
        calls.append(Path(path))
        return real_read_pdf(path)

    monkeypatch.setattr(specialists, "read_pdf", spy_read_pdf)
    stub_fields = [{"label": "Study ID", "collected_variable": None}]
    inst = _instrument(reply={"fields": stub_fields})
    result = asyncio.run(inst.run(form_files=[_form_file("f1", FLAT_PDF)]))

    assert calls == [FLAT_PDF]
    assert len(inst.call_json_calls) == 1
    assert result["fields"] == stub_fields
    assert inst._fields["f1"] == stub_fields
    # The real digital text layer reached the prompt (scrubbed), proving no
    # inline pypdf re-parse replaced the shared reader's output.
    assert "Study ID" in inst.call_json_calls[0]["user_prompt"]


def test_no_inline_pdf_parsing_left_in_specialists_module():
    """Regression for the removed inline pypdf.PdfReader Tier-2 read."""
    assert not hasattr(specialists, "PdfReader")


# ---- Task 12: scanned/OCR-dependent PDF routes through read_pdf, and its
#      output is scrubbed before the LLM ever sees it -----------------------


def test_scanned_form_routes_through_shared_read_pdf_and_scrubs_before_prompt(monkeypatch):
    calls: list[Path] = []
    phi_text = "Patient Name: James Smith, 415-555-1234\nSite of Disease: pulmonary"

    def stub_read_pdf(path):
        calls.append(Path(path))
        return phi_text

    monkeypatch.setattr(specialists, "read_pdf", stub_read_pdf)
    inst = _instrument(reply={"fields": [{"label": "Site of Disease", "collected_variable": None}]})
    asyncio.run(inst.run(form_files=[_form_file("f1", SCANNED_PDF)]))

    assert calls == [SCANNED_PDF]
    assert len(inst.call_json_calls) == 1
    prompt = inst.call_json_calls[0]["user_prompt"]
    # detectors=("rule",) catches structurally-shaped PHI (phone numbers,
    # SSNs, dates, emails); it deliberately does not run NER over free
    # text, so a bare name is Judge/Sentinel's downstream concern, not
    # Instrument's prompt-scrub responsibility -- this is the same
    # rule-only tradeoff every other agent's call() makes.
    assert "415-555-1234" not in prompt
    assert "[REDACTED:D:PHONE]" in prompt


def _ocr_stack_available() -> bool:
    try:
        import pytesseract  # noqa: F401
        import pdf2image  # noqa: F401
    except Exception:
        return False
    return bool(which("tesseract") and which("pdftoppm"))


@pytest.mark.skipif(not _ocr_stack_available(),
                    reason="tesseract / pdf2image / poppler not available in this environment")
def test_scanned_fixture_ocr_extracts_the_baked_in_text():
    from phi_core.file_readers import read_pdf, OCR_TEXT_THRESHOLD

    extracted = read_pdf(SCANNED_PDF)
    assert len(extracted.strip()) >= OCR_TEXT_THRESHOLD
    joined = extracted.lower()
    assert "james" in joined or "smith" in joined
    assert "555" in joined and "1234" in joined


# ---- Task 12: .docx form routes through _read_docx_tables + read_docx ------


def _write_docx_form(path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("TB collection form, Word version.")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Variable"
    table.rows[1].cells[0].text = "Study ID"
    table.rows[1].cells[1].text = "study_id"
    doc.save(str(path))


def test_docx_form_combines_table_and_prose_text(tmp_path, monkeypatch):
    docx_path = tmp_path / "tb_form.docx"
    _write_docx_form(docx_path)

    inst = _instrument(reply={"fields": [{"label": "Study ID", "collected_variable": "study_id"}]})
    asyncio.run(inst.run(form_files=[_form_file("f1", docx_path)]))

    assert len(inst.call_json_calls) == 1
    prompt = inst.call_json_calls[0]["user_prompt"]
    assert "Word version" in prompt          # from file_readers.read_docx (prose)
    assert "study_id" in prompt              # from _read_docx_tables (table cell)


# ---- verify(): deterministic, no LLM call -----------------------------------


def test_verify_matches_label_or_variable_case_insensitive():
    inst = _instrument(NoLlmInstrument)
    asyncio.run(inst.run(form_files=[_form_file("f1", ACROFORM_PDF)]))

    by_var = inst.verify("STUDY_ID")
    assert by_var["present"] is True
    assert by_var["file_id"] == "f1"
    assert by_var["field"]["collected_variable"] == "study_id"

    by_label = inst.verify(inst._fields["f1"][0]["label"].upper())
    assert by_label["present"] is True

    absent = inst.verify("not_a_real_field")
    assert absent == {
        "present": False,
        "explanation": "not present in any extracted form field",
    }

    # scoping to a file_id that doesn't have the field
    scoped_absent = inst.verify("study_id", file_id="nonexistent_file")
    assert scoped_absent["present"] is False


# ---- Task 13: per-form report artifact --------------------------------------


def test_report_written_from_in_memory_fields_with_no_phi_category(tmp_path, monkeypatch):
    monkeypatch.setattr(paths_mod, "UPLOAD_DIR", tmp_path)
    inst = _instrument(NoLlmInstrument)
    asyncio.run(inst.run(form_files=[_form_file("f1", ACROFORM_PDF, "tb_collection_form_acroform.pdf")]))

    report_path = tmp_path / "s1" / "instrument_report_f1.json"
    assert report_path.exists()
    payload = json.loads(report_path.read_text(encoding="utf-8"))
    assert payload["file_id"] == "f1"
    assert payload["source_filename"] == "tb_collection_form_acroform.pdf"
    in_memory = inst._fields["f1"]
    assert len(payload["fields"]) == len(in_memory)
    # scrub_persisted_text is intentionally over-cautious about Title-Case
    # label text (a filled/OCR'd form can put a real name into a label),
    # so labels may be redacted; collected_variable is a snake_case
    # identifier the name detector never touches, and a plain
    # single-word label (no PHI shape) round-trips verbatim.
    for written, original in zip(payload["fields"], in_memory):
        assert written["collected_variable"] == original["collected_variable"]
    written_by_var = {f["collected_variable"]: f["label"] for f in payload["fields"]}
    assert written_by_var["city"] == "City:"
    assert written_by_var["zip_code"] == "ZIP Code [zip_code]:"
    assert "phi_category" not in json.dumps(payload)


def test_report_is_written_from_fields_not_agent_log(tmp_path, monkeypatch):
    """The report must come from self._fields, never a reconstruction of
    agent_log -- Agent._log's write-time scrub mangles free-text labels."""
    monkeypatch.setattr(paths_mod, "UPLOAD_DIR", tmp_path)
    inst = _instrument(reply={"fields": [{"label": "Weird [REDACTED:A:NAME] Label", "collected_variable": None}]})
    asyncio.run(inst.run(form_files=[_form_file("f1", FLAT_PDF)]))

    report_path = tmp_path / "s1" / "instrument_report_f1.json"
    payload = json.loads(report_path.read_text(encoding="utf-8"))
    # Written straight from self._fields: the label round-trips unscrubbed
    # by the trace-layer scrubber (scrub_persisted_text may still touch it,
    # but never through agent_log's reply_text mangling).
    assert payload["fields"][0]["label"] == inst._fields["f1"][0]["label"]
