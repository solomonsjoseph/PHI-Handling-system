"""Real-world file-shape regression tests.

Sir uploaded three example artefacts on 2026-02 (Word dictionary, CSV
dataset, annotated CRF PDF) with the directive "handling of these types
of files is must, the dictionary is word doc here but if large then it
would be excel workbook. Align it accordingly." These tests lock the
.docx dictionary parser + xls fallback so a future refactor cannot drop
them.
"""
from __future__ import annotations

import zipfile
from pathlib import Path


DOCX_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    <w:p><w:r><w:t>Diabetes Care QI Data Dictionary</w:t></w:r></w:p>
    <w:tbl>
      <w:tr>
        <w:tc><w:p><w:r><w:t>Column</w:t></w:r></w:p></w:tc>
        <w:tc><w:p><w:r><w:t>Type</w:t></w:r></w:p></w:tc>
        <w:tc><w:p><w:r><w:t>What it represents</w:t></w:r></w:p></w:tc>
      </w:tr>
      <w:tr>
        <w:tc><w:p><w:r><w:t>patient_id</w:t></w:r></w:p></w:tc>
        <w:tc><w:p><w:r><w:t>ID</w:t></w:r></w:p></w:tc>
        <w:tc><w:p><w:r><w:t>Unique patient identifier</w:t></w:r></w:p></w:tc>
      </w:tr>
      <w:tr>
        <w:tc><w:p><w:r><w:t>encounter_date</w:t></w:r></w:p></w:tc>
        <w:tc><w:p><w:r><w:t>Date</w:t></w:r></w:p></w:tc>
        <w:tc><w:p><w:r><w:t>Visit date, ISO format</w:t></w:r></w:p></w:tc>
      </w:tr>
    </w:tbl>
    <w:p><w:r><w:t>End of dictionary.</w:t></w:r></w:p>
  </w:body>
</w:document>"""


def _make_docx(path: Path, xml: str = DOCX_XML) -> None:
    with zipfile.ZipFile(path, "w") as z:
        z.writestr("word/document.xml", xml)


def test_read_docx_tables_extracts_word_table(tmp_path):
    from phi_core.agents.specialists import _read_docx_tables
    p = tmp_path / "dict.docx"
    _make_docx(p)
    out = _read_docx_tables(p)
    assert "# table 1" in out
    assert "Column,Type,What it represents" in out
    assert "patient_id,ID,Unique patient identifier" in out
    assert "encounter_date,Date," in out
    # Cell value with a comma must be quoted correctly
    assert '"Visit date, ISO format"' in out


def test_read_docx_tables_captures_narrative_context(tmp_path):
    from phi_core.agents.specialists import _read_docx_tables
    p = tmp_path / "dict.docx"
    _make_docx(p)
    out = _read_docx_tables(p)
    assert "# narrative context" in out
    assert "Diabetes Care QI Data Dictionary" in out
    assert "End of dictionary." in out


def test_read_docx_tables_returns_empty_on_bad_zip(tmp_path):
    from phi_core.agents.specialists import _read_docx_tables
    p = tmp_path / "not-a-docx.docx"
    p.write_bytes(b"this is not a docx")
    assert _read_docx_tables(p) == ""


# ---- SEC-001 iter_20: nested-docx decompression bomb -------------------


def test_read_docx_tables_rejects_oversize_document_xml(tmp_path):
    """A malicious docx can pack a >10 MiB document.xml into a tiny
    compressed archive. The reader must refuse it rather than allocate
    unbounded RAM in ET.parse. Return "" so Lexicon just sees no
    dictionary text; no crash, no OOM."""
    from phi_core.agents.specialists import _read_docx_tables, _DOCX_XML_MAX_BYTES
    import zipfile as _zip
    p = tmp_path / "bomb.docx"
    with _zip.ZipFile(p, "w", _zip.ZIP_DEFLATED) as z:
        # Compressible payload > cap
        big = (b"<w:document xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'>"
               b"<w:body>" + b" " * (_DOCX_XML_MAX_BYTES + 500_000) + b"</w:body></w:document>")
        z.writestr("word/document.xml", big)
    assert p.stat().st_size < 200_000, "test setup: outer archive should be small"
    assert _read_docx_tables(p) == "", "oversized inner xml must be rejected"


# ---- SEC-002 iter_20: billion-laughs DTD refusal ------------------------


def test_read_docx_tables_refuses_dtd_via_defusedxml(tmp_path):
    """defusedxml with forbid_dtd=True must refuse ANY <!DOCTYPE ...>
    declaration. Blocks billion-laughs / quadratic-blowup entity
    expansion even on hosts with older libexpat that lack the built-in
    amplification protection."""
    from phi_core.agents.specialists import _read_docx_tables
    import zipfile as _zip
    p = tmp_path / "dtd.docx"
    with _zip.ZipFile(p, "w") as z:
        z.writestr("word/document.xml", (
            b"<?xml version='1.0'?>"
            b"<!DOCTYPE lolz [<!ENTITY lol 'lol'>]>"
            b"<w:document xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'>"
            b"<w:body><w:p><w:r><w:t>&lol;</w:t></w:r></w:p></w:body></w:document>"
        ))
    assert _read_docx_tables(p) == "", "DTD-bearing docx must be refused"


# ---- SEC-001 iter_22: narrative-path docx reader parity ---------------


def test_narrative_read_docx_rejects_oversize_document_xml(tmp_path):
    """The narrative-file reader (`file_readers.read_docx`) must apply
    the same size cap as the dictionary reader. Both readers now share
    `phi_core.docx_safe.DOCX_XML_MAX_BYTES` so a future adjustment applies
    to both paths at once (iter_23 shared-helper refactor)."""
    from phi_core.file_readers import read_docx
    from phi_core.docx_safe import DOCX_XML_MAX_BYTES
    import zipfile as _zip
    p = tmp_path / "bomb.docx"
    with _zip.ZipFile(p, "w", _zip.ZIP_DEFLATED) as z:
        big = (b"<?xml version='1.0'?>"
               b"<w:document xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main'>"
               b"<w:body>" + b" " * (DOCX_XML_MAX_BYTES + 500_000) + b"</w:body></w:document>")
        z.writestr("word/document.xml", big)
    assert p.stat().st_size < 200_000
    assert read_docx(p) == "", "oversized inner xml on narrative path must be refused"


def test_docx_safe_helper_is_single_source_of_truth():
    """iter_23: both docx readers must consult `phi_core.docx_safe` so
    the bomb + DTD defence cannot drift again."""
    import inspect
    from phi_core.file_readers import read_docx
    from phi_core.agents.specialists import _read_docx_tables
    assert "docx_safe" in inspect.getsource(read_docx)
    assert "docx_safe" in inspect.getsource(_read_docx_tables)


def test_narrative_read_docx_returns_empty_on_bad_zip(tmp_path):
    """Malformed / non-zip .docx must not raise; caller must see ""."""
    from phi_core.file_readers import read_docx
    p = tmp_path / "junk.docx"
    p.write_bytes(b"this is not a docx")
    assert read_docx(p) == ""


def test_narrative_read_docx_still_reads_normal_docx(tmp_path):
    """Positive path: a legitimate small .docx returns its paragraph
    text unchanged after the iter_22 cap was added."""
    from docx import Document
    from phi_core.file_readers import read_docx
    p = tmp_path / "consent.docx"
    d = Document()
    d.add_paragraph("Patient consent form")
    d.add_paragraph("Signed 2024-05-20")
    d.save(str(p))
    out = read_docx(p)
    assert "Patient consent form" in out
    assert "Signed 2024-05-20" in out


def test_read_table_flat_dispatches_docx(tmp_path):
    from phi_core.agents.specialists import _read_table_flat
    p = tmp_path / "dict.docx"
    _make_docx(p)
    out = _read_table_flat(p)
    assert "patient_id" in out


def test_intake_accepts_docx_dictionary():
    from phi_core.intake import COMPONENT_SUFFIXES
    assert ".docx" in COMPONENT_SUFFIXES["dictionary"]
    # xls stays supported too (Sir: "if large then it would be excel workbook")
    assert ".xlsx" in COMPONENT_SUFFIXES["dictionary"]
    assert ".xls" in COMPONENT_SUFFIXES["dictionary"]


def test_read_xls_tables_handles_xlsx_disguised_as_xls(tmp_path):
    """Some data stewards rename .xlsx to .xls. The reader must NOT
    crash regardless of what actually landed on disk. Without xlrd
    installed and with openpyxl refusing the .xls suffix, the reader
    returns an empty string (Lexicon just receives no table text)."""
    from openpyxl import Workbook
    from phi_core.agents.specialists import _read_xls_tables
    p_xlsx = tmp_path / "dict.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append(["Column", "Type", "What it represents"])
    ws.append(["patient_id", "ID", "Unique patient identifier"])
    wb.save(str(p_xlsx))
    p_xls = tmp_path / "dict.xls"
    p_xls.write_bytes(p_xlsx.read_bytes())
    # Must not raise. xlrd may or may not be installed; either way, the
    # function is contract-bound to return a string, never None.
    out = _read_xls_tables(p_xls)
    assert isinstance(out, str)


def test_metadata_xlsx_scrubs_every_sheet_and_intake_rejects_multisheet_dictionary(tmp_path):
    """Dictionary workbooks must never retain PHI on an unscanned worksheet."""
    from openpyxl import Workbook, load_workbook
    from phi_core.agents.reasoning import _redact_metadata_file
    from phi_core.intake import scan_intake

    src = tmp_path / "dictionary" / "codebook.xlsx"
    src.parent.mkdir()
    wb = Workbook()
    wb.active.append(["description"])
    wb.active.append(["Callback phone 415-555-1234"])
    second = wb.create_sheet("supplement")
    second.append(["description"])
    second.append(["Contact james@example.edu"])
    wb.save(src)

    redacted = _redact_metadata_file(src, tmp_path / "redacted.xlsx")
    redacted_wb = load_workbook(redacted)
    values = [
        str(cell.value)
        for worksheet in redacted_wb.worksheets
        for row in worksheet.iter_rows()
        for cell in row
        if cell.value is not None
    ]
    assert not any("415-555-1234" in value for value in values)
    assert not any("james@example.edu" in value for value in values)

    datasets = tmp_path / "datasets"
    datasets.mkdir()
    (datasets / "study.csv").write_text("study_id\n1\n", encoding="utf-8")
    entries, _ = scan_intake(tmp_path)
    rejected = [entry for entry in entries if entry.relpath == "dictionary/codebook.xlsx"]
    assert len(rejected) == 1
    assert rejected[0].component == "_unclassified"
    assert "xlsx has 2 sheets" in rejected[0].reason


def test_withheld_metadata_uses_scannable_txt_destination(tmp_path):
    """Unsupported dictionary files must be withheld at a guard-scannable path."""
    from phi_core.agents.reasoning import _redact_metadata_file
    from phi_core.publish_guard import scan_export_file

    src = tmp_path / "codebook.docx"
    src.write_text("Jane Q. Patient, MRN 4471129", encoding="utf-8")
    actual = _redact_metadata_file(src, tmp_path / "codebook.docx")

    assert actual == tmp_path / "codebook.withheld.txt"
    assert actual.exists()
    assert "Jane Q. Patient" not in actual.read_text(encoding="utf-8")
    assert scan_export_file("codebook", actual).status == "clean"


def test_read_export_rows_includes_every_xlsx_worksheet(tmp_path):
    """Verifier input must include rows from sheets after the first."""
    from openpyxl import Workbook
    from phi_corpus.verify import _read_export_rows

    export = tmp_path / "export.xlsx"
    wb = Workbook()
    wb.active.append(["first"])
    wb.active.append(["one"])
    second = wb.create_sheet("supplement")
    second.append(["second"])
    second.append(["two"])
    wb.save(export)

    assert _read_export_rows(str(export)) == [["first"], ["one"], ["second"], ["two"]]


def test_intake_accepted_xls_metadata_is_withheld_at_scannable_path(tmp_path):
    """Legacy dictionary XLS files must be withheld without openpyxl parsing."""
    from phi_core.agents.reasoning import _redact_metadata_file
    from phi_core.intake import scan_intake

    src = tmp_path / "dictionary" / "legacy.xls"
    src.parent.mkdir()
    src.write_text("Jane Q. Patient, MRN 4471129", encoding="utf-8")
    datasets = tmp_path / "datasets"
    datasets.mkdir()
    (datasets / "study.csv").write_text("study_id\n1\n", encoding="utf-8")

    entries, _ = scan_intake(tmp_path)
    accepted = [entry for entry in entries if entry.relpath == "dictionary/legacy.xls"]
    assert len(accepted) == 1
    assert accepted[0].component == "dictionary"

    actual = _redact_metadata_file(src, tmp_path / "legacy.xls")
    assert actual == tmp_path / "legacy.withheld.txt"
    assert actual.read_text(encoding="utf-8").startswith("[REDACTED]")


def test_executor_publishes_withheld_metadata_at_its_scannable_path(tmp_path, monkeypatch):
    """Executor exports the actual withheld marker path for unsupported metadata."""
    import asyncio

    from phi_core.agents.llm import LlmConfig
    from phi_core.agents.reasoning import Executor
    from phi_core.publish_guard import scan_export_file
    import phi_core.agents.reasoning as reasoning

    class _AgentLog:
        async def insert_one(self, _message):
            return None

    class _Database:
        agent_log = _AgentLog()

    src = tmp_path / "codebook.docx"
    src.write_text("Jane Q. Patient, MRN 4471129", encoding="utf-8")
    export_dir = tmp_path / "exports"
    export_dir.mkdir()
    monkeypatch.setattr(reasoning, "EXPORT_DIR", export_dir)

    result = asyncio.run(Executor("session", LlmConfig(), _Database()).run(
        [{
            "file_id": "dictionary",
            "stored_path": str(src),
            "original_name": src.name,
            "kind": "metadata",
        }],
        [],
    ))

    actual = export_dir / "dictionary__codebook.withheld.txt"
    assert result["exports"]["dictionary"] == str(actual)
    assert scan_export_file("dictionary", actual).status == "clean"
