"""Phase 14 (scale and resilience): multi-dataset studies at scale.

Section 97 asks for coverage over "many columns, many forms, many
dictionaries, duplicate column names, multi-sheet files where supported".
This file scales Phase 7's coverage-invariant fixture
(``test_control_phase7_coverage_invariant.py``, two files, one duplicate
column name each) up by an order of magnitude -- many files, many columns
per file, and heavy duplicate-name collision across files -- and proves
the (file_id, column) identity invariant still holds exactly at that
scale, then does the same for Lexicon (many dictionary files), Schema
(many dataset files), and Instrument (many form files), plus a
multi-sheet xlsx read.

No production code is modified; every test below exercises real,
unmodified production functions (``gates.assert_exact_coverage``,
``gates.run_decision_gates``, ``Lexicon.run``, ``Schema.run``,
``Instrument.run``, ``file_readers.read_xlsx_columns`` /
``iter_dataset_rows``).
"""
from __future__ import annotations

import csv
import io
from pathlib import Path
from uuid import uuid4

import openpyxl
import pytest
from phi_core.agents.specialists import Instrument, Lexicon, Schema
from phi_core.control.gates import assert_exact_coverage, run_decision_gates
from phi_core.control.records import WorkflowRun
from phi_core.control.store import MemoryControlStore
from phi_core.control.testing import FakeGateway, make_ctx
from phi_core.file_readers import iter_dataset_rows, read_xlsx_columns
from phi_core.paths import DATA_DIR

_N_FILES = 12
_COLS_PER_FILE = 15
# Every file shares this exact column name with every other file --
# Phase 7's own "two files literally named 'notes'" collision, widened
# to N files instead of 2, to prove the (file_id, column) identity holds
# under heavier collision, not just a single pair.
_SHARED_NAME = "notes"


def _decision(file_id: str, column: str, **overrides: object) -> dict[str, object]:
    base = {
        "file_id": file_id, "column": column, "action": "keep", "reason": "scale fixture",
        "confidence": 0.9, "phi_category": "NONE", "subject": "participant", "citation": "",
    }
    base.update(overrides)
    return base


def _many_files() -> list[dict[str, object]]:
    """``_N_FILES`` files, each with ``_COLS_PER_FILE`` columns. Every
    file's last column is literally named ``_SHARED_NAME`` -- the same
    name repeated across all ``_N_FILES`` files, distinguishable only by
    ``file_id`` per docs section 31 (Judge's classification identity is
    ``(file_id, column_id)``, never the bare name)."""
    files = []
    for fi in range(_N_FILES):
        file_id = f"f{fi}"
        columns = [f"col_{fi}_{ci}" for ci in range(_COLS_PER_FILE - 1)] + [_SHARED_NAME]
        files.append({"file_id": file_id, "columns": columns, "stored_path": ""})
    return files


def _decisions_for(files: list[dict[str, object]]) -> list[dict[str, object]]:
    return [_decision(f["file_id"], column) for f in files for column in f["columns"]]


# ---- coverage invariant at scale -------------------------------------------


def test_assert_exact_coverage_holds_over_many_files_with_heavy_duplicate_names() -> None:
    files = _many_files()
    decisions = _decisions_for(files)

    status, detail = assert_exact_coverage(decisions, files)

    assert status == "pass", detail
    total_columns = _N_FILES * _COLS_PER_FILE
    assert f"exact coverage over {total_columns} column(s) across {_N_FILES} file(s)" in detail
    # Every one of the _N_FILES "notes" columns is a distinct logical
    # column: (file_id, "notes") pairs are exactly _N_FILES-many, never
    # collapsed to 1 by the shared bare name.
    shared_pairs = {(d["file_id"], d["column"]) for d in decisions if d["column"] == _SHARED_NAME}
    assert len(shared_pairs) == _N_FILES


@pytest.mark.asyncio
async def test_run_decision_gates_gives_every_logical_column_exactly_one_decision_at_scale() -> None:
    files = _many_files()
    decisions = _decisions_for(files)
    store = MemoryControlStore()
    run_id = uuid4().hex
    await store.insert("workflow_runs", WorkflowRun(
        run_id=run_id, session_id=run_id, workflow_version="wf/1", policy_version="p/1",
        state="running", node="decide",
        checkpoint={"node": "decide", "checkpoint_version": 1, "payload_refs": []},
    ))
    ctx = make_ctx("Judge", run_id=run_id, store=store)

    outcome = await run_decision_gates(
        decisions=decisions, files=files, stage="decide", ctx=ctx, store=store,
    )

    assert outcome.ok is True
    assert len(outcome.decisions) == _N_FILES * _COLS_PER_FILE


@pytest.mark.asyncio
async def test_run_decision_gates_fails_closed_on_one_missing_decision_buried_among_many() -> None:
    """One missing decision -- out of ``_N_FILES * _COLS_PER_FILE`` -- is
    still caught precisely, at scale: the failure names the exact
    (file_id, column) pair, not a vague "coverage incomplete"."""
    files = _many_files()
    decisions = _decisions_for(files)
    # Drop exactly one decision, buried in the middle of the list: file
    # "f6"'s "notes" column (one of the _N_FILES colliding-name entries).
    victim = ("f6", _SHARED_NAME)
    decisions = [d for d in decisions if (d["file_id"], d["column"]) != victim]
    store = MemoryControlStore()
    run_id = uuid4().hex
    await store.insert("workflow_runs", WorkflowRun(
        run_id=run_id, session_id=run_id, workflow_version="wf/1", policy_version="p/1",
        state="running", node="decide",
        checkpoint={"node": "decide", "checkpoint_version": 1, "payload_refs": []},
    ))
    ctx = make_ctx("Judge", run_id=run_id, store=store)

    outcome = await run_decision_gates(
        decisions=decisions, files=files, stage="decide", ctx=ctx, store=store,
    )

    assert outcome.ok is False
    coverage_result = next(gr for gr in outcome.gate_results if gr.gate == "assert_exact_coverage")
    assert f"missing_decision:[{victim!r}]" in coverage_result.detail


# ---- multi-sheet xlsx -------------------------------------------------------


def _write_multi_sheet_xlsx(path: Path) -> None:
    wb = openpyxl.Workbook()
    sheet1 = wb.active
    sheet1.title = "dataset"
    sheet1.append(["subject_id", "age", "zip"])
    sheet1.append(["p1", "34", "021"])
    sheet1.append(["p2", "51", "029"])

    sheet2 = wb.create_sheet("codebook")
    sheet2.append(["totally_different_columns", "should_never_appear"])
    sheet2.append(["x", "y"])

    sheet3 = wb.create_sheet("notes")
    sheet3.append(["another_unrelated_sheet"])

    wb.save(path)


def test_multi_sheet_xlsx_reads_the_first_sheet_deterministically_and_ignores_the_rest() -> None:
    path = DATA_DIR / "uploads" / f"{uuid4().hex}.xlsx"
    _write_multi_sheet_xlsx(path)

    columns, row_count = read_xlsx_columns(path)

    assert columns == ["subject_id", "age", "zip"]
    assert row_count == 2
    assert "totally_different_columns" not in columns
    assert "another_unrelated_sheet" not in columns


def test_iter_dataset_rows_over_a_multi_sheet_workbook_yields_only_the_first_sheets_rows() -> None:
    path = DATA_DIR / "uploads" / f"{uuid4().hex}.xlsx"
    _write_multi_sheet_xlsx(path)

    rows = list(iter_dataset_rows(path, "xlsx"))

    assert len(rows) == 2
    _idx0, values0 = rows[0]
    assert set(values0) == {"subject_id", "age", "zip"}
    assert values0["subject_id"] == "p1"
    assert all("totally_different_columns" not in v for _i, v in rows)


# ---- many dictionaries (Lexicon) -------------------------------------------

_N_DICT_FILES = 10
_ROWS_PER_DICT = 25


def _write_dict_csv(path: Path, file_index: int) -> list[str]:
    names = [f"dict{file_index}_var{i}" for i in range(_ROWS_PER_DICT)]
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(["variable_name", "description"])
    for name in names:
        writer.writerow([name, f"description of {name}"])
    path.write_text(buf.getvalue(), encoding="utf-8")
    return names


@pytest.mark.asyncio
async def test_lexicon_indexes_every_row_across_many_dictionary_files_even_with_llm_outage() -> None:
    """``_N_DICT_FILES`` dictionary files, ``_ROWS_PER_DICT`` rows each
    (``_N_DICT_FILES * _ROWS_PER_DICT`` total documented columns). The
    gateway is a ``FakeGateway`` with an empty replies queue -- every
    ``call_json`` gets the empty-string default, simulating a total LLM
    outage during the gist-filling pass. Lexicon's own docstring
    guarantee (deterministic row extraction happens before any LLM call)
    means every row must still appear in the output, with a blank gist
    rather than a dropped row."""
    dict_files = []
    all_names: list[str] = []
    for fi in range(_N_DICT_FILES):
        path = DATA_DIR / "uploads" / f"{uuid4().hex}.csv"
        names = _write_dict_csv(path, fi)
        all_names.extend(names)
        dict_files.append({"file_id": f"d{fi}", "stored_path": str(path)})

    gateway = FakeGateway()  # empty replies deque: every call_json gets "" -> default
    ctx = make_ctx("Lexicon", gateway=gateway)
    lexicon = Lexicon(ctx)

    result = await lexicon.run(dict_files=dict_files)

    returned_names = {c["name"] for c in result["columns"]}
    assert returned_names == set(all_names)
    assert len(result["columns"]) == _N_DICT_FILES * _ROWS_PER_DICT
    # LLM outage: every gist is blank, but no row was dropped.
    assert all(c["description"] == "" for c in result["columns"])


# ---- many dataset files, many columns (Schema) -----------------------------

_N_SCHEMA_FILES = 20
_SCHEMA_COLS_PER_FILE = 30


def _write_schema_csv(path: Path, file_index: int) -> list[str]:
    columns = [f"s{file_index}_c{i}" for i in range(_SCHEMA_COLS_PER_FILE)]
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow(columns)
    writer.writerow([str(i) for i in range(_SCHEMA_COLS_PER_FILE)])
    writer.writerow([str(i * 2) for i in range(_SCHEMA_COLS_PER_FILE)])
    path.write_text(buf.getvalue(), encoding="utf-8")
    return columns


@pytest.mark.asyncio
async def test_schema_reads_many_dataset_files_with_many_columns_without_dropping_any() -> None:
    from phi_core.agents.extract_model import ExtractedColumn, ExtractedSchema
    from phi_core.file_readers import read_csv_columns

    class _FileReadingSchema(Schema):
        """Stubs only the codegen/LLM/Docker extraction seam (step 10);
        every other line of ``Schema.run`` -- the per-file grouping this
        test actually pins -- is the real, unmodified production
        method, matching this file's own module-level intent."""

        async def _extract_via_codegen(self, f, sandbox):
            headers, _rows = read_csv_columns(Path(f["stored_path"]))
            columns = [
                ExtractedColumn(name=h, position=i, distinct_count=2, null_count=0, inferred_type="string")
                for i, h in enumerate(headers)
            ]
            return ExtractedSchema(columns=columns, row_count=10)

    dataset_files = []
    expected: dict[str, list[str]] = {}
    for fi in range(_N_SCHEMA_FILES):
        path = DATA_DIR / "uploads" / f"{uuid4().hex}.csv"
        columns = _write_schema_csv(path, fi)
        file_id = f"sf{fi}"
        expected[file_id] = columns
        dataset_files.append({"file_id": file_id, "stored_path": str(path), "subtype": "csv"})

    schema = _FileReadingSchema(make_ctx("Schema"))
    result = await schema.run(dataset_files=dataset_files)

    by_file: dict[str, list[str]] = {}
    for r in result["columns"]:
        by_file.setdefault(r["_file_id"], []).append(r["name"])
    assert set(by_file) == set(expected)
    for file_id, columns in expected.items():
        assert by_file[file_id] == columns


# ---- many forms (Instrument) ------------------------------------------------

_N_FORM_FILES = 8


def _write_form_docx(path: Path, field_labels: list[str]) -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Clinical study intake form")
    for label in field_labels:
        doc.add_paragraph(f"{label}: ____________")
    doc.save(str(path))


@pytest.mark.asyncio
async def test_instrument_processes_many_form_files_without_dropping_any_file() -> None:
    form_files = []
    for fi in range(_N_FORM_FILES):
        path = DATA_DIR / "uploads" / f"{uuid4().hex}.docx"
        labels = [f"field_{fi}_{j}" for j in range(3)]
        _write_form_docx(path, labels)
        form_files.append({"file_id": f"form{fi}", "stored_path": str(path), "subtype": "docx"})

    gateway = FakeGateway()
    for _ in form_files:
        gateway.replies.append('{"fields": [{"label": "some field", "collected_variable": null}]}')
    ctx = make_ctx("Instrument", gateway=gateway)
    instrument = Instrument(ctx)

    result = await instrument.run(form_files=form_files)

    # Every form file was actually sent to the gateway (one call per
    # file that reached Tier 2, since none of these .docx files are
    # AcroForm PDFs) -- no file silently skipped at this scale.
    assert len(gateway.requests) == _N_FORM_FILES
    assert set(instrument._fields) == {f["file_id"] for f in form_files}
    assert len(result["fields"]) == _N_FORM_FILES  # one fixed reply's field, per file
